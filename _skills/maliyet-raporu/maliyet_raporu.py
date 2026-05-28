# ═══════════════════════════════════════════════════════
# ANTIGRAVITY — Maliyet Raporu Üretici
# Kullanım: python maliyet_raporu.py --proje "PROJE_ADI" --klasor "C:\yol\klasor"
# ═══════════════════════════════════════════════════════
import argparse
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def maliyet_raporu_olustur(
    proje_adi: str,
    klasor: str,
    satirlar: list[dict],          # [{"aciklama": ..., "kredi": ..., "maliyet": ...}, ...]
    diger_maliyetler: list[dict],  # [{"servis": ..., "maliyet": ...}, ...]
    toplam_usd: float,
    tarih: str = None,
    kredi_birimi: str = "25 kredi/saniye  |  1000 kredi = 5 USD"
):
    """
    Standart Antigravity maliyet raporu oluşturur ve proje klasörüne kaydeder.

    Args:
        proje_adi: Proje adı (örn. "LARA ARI")
        klasor: Kayıt klasörü tam yolu
        satirlar: Ana tablo satırları
        diger_maliyetler: Diğer servis maliyetleri
        toplam_usd: Genel toplam (float)
        tarih: Tarih string (varsayılan: bugün)
        kredi_birimi: Kredi açıklaması
    """
    if tarih is None:
        tarih = datetime.now().strftime("%d %B %Y")

    doc = Document()

    # Sayfa marginleri
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Ana başlık
    h = doc.add_heading('', level=0)
    run = h.add_run(f'{proje_adi} - Video Uretim Maliyet Raporu')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Alt başlık
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p.add_run(f'Tarih: {tarih}  |  Proje: {proje_adi} UGC Reklam Filmi')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Ana maliyet tablosu başlığı
    h2 = doc.add_heading('', level=1)
    r2 = h2.add_run('Kie AI (Seedance 2.0) - Video Uretim')
    r2.font.size = Pt(13)
    r2.font.bold = True

    info = doc.add_paragraph()
    info.add_run(f'Birim fiyat: {kredi_birimi}').font.size = Pt(9)

    # Tablo
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    basliklar = ['Video / Aciklama', 'Uretilen Icerik', 'Kullanilan Kredi', 'Maliyet (USD)']
    for i, b in enumerate(basliklar):
        hdr = table.rows[0].cells[i]
        hdr.text = b
        hdr.paragraphs[0].runs[0].bold = True
        hdr.paragraphs[0].runs[0].font.size = Pt(9)

    for satir in satirlar:
        row = table.add_row().cells
        degerler = [
            satir.get('video', ''),
            satir.get('aciklama', ''),
            satir.get('kredi', ''),
            satir.get('maliyet', '')
        ]
        for i, val in enumerate(degerler):
            row[i].text = val
            row[i].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # Diğer servisler
    h3 = doc.add_heading('', level=1)
    r3 = h3.add_run('Diger Servisler')
    r3.font.size = Pt(13)
    r3.font.bold = True

    for item in diger_maliyetler:
        p2 = doc.add_paragraph(style='List Bullet')
        p2.add_run(item['servis'] + ':  ').font.size = Pt(10)
        br = p2.add_run(item['maliyet'])
        br.bold = True
        br.font.size = Pt(10)

    doc.add_paragraph()

    # TOPLAM
    toplam_p = doc.add_paragraph()
    toplam_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = toplam_p.add_run(f'GENEL TOPLAM:  {toplam_usd:.2f} USD'.replace('.', ','))
    tr.font.size = Pt(16)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)

    doc.add_paragraph()

    # Not
    not_p = doc.add_paragraph()
    nr = not_p.add_run(
        'NOT: Bu rapor yalnizca video uretim API maliyetlerini kapsamaktadir. '
        'Platform abonelik ucretleri bu rapora dahil edilmemistir.'
    )
    nr.font.size = Pt(8)
    nr.font.italic = True
    nr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f'Rapor: Antigravity Otomasyon Sistemi  |  {tarih}')
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # Kaydet
    dosya_adi = f'{proje_adi.replace(" ", "_")}_Maliyet_Raporu.docx'
    tam_yol = os.path.join(klasor, dosya_adi)
    os.makedirs(klasor, exist_ok=True)
    doc.save(tam_yol)
    print(f'Rapor kaydedildi: {tam_yol}')
    return tam_yol


# ── Örnek kullanım (direkt çalıştırma için) ──
if __name__ == "__main__":
    ornek_satirlar = [
        {"video": "VIDEO_01.mp4", "aciklama": "3 sahne x 15 saniye", "kredi": "375 kredi", "maliyet": "1,88 USD"},
        {"video": "VIDEO_02.mp4", "aciklama": "FFmpeg birlestirme",   "kredi": "0 kredi",   "maliyet": "0,00 USD"},
    ]
    ornek_diger = [
        {"servis": "ElevenLabs TTS", "maliyet": "0,08 USD"},
        {"servis": "Replicate",      "maliyet": "0,06 USD"},
    ]
    maliyet_raporu_olustur(
        proje_adi="ORNEK_PROJE",
        klasor=r"C:\Users\msist\OneDrive\Desktop\Antigravity(DOLUNAY)\hlk-REKLAM",
        satirlar=ornek_satirlar,
        diger_maliyetler=ornek_diger,
        toplam_usd=2.02
    )
